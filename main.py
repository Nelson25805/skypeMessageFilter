import json
import html
import re
import os
from docx import Document
from datetime import datetime

def load_conversation(json_file):
    """Load the Skype conversation JSON file."""
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data

def clean_content(text):
    """
    Clean the message content by:
    1. Unescaping HTML entities.
    2. Removing XML/HTML tags like <e_m ...></e_m>.
    """
    text = html.unescape(text)
    text = re.sub(r'<e_m\b[^>]*>(.*?)</e_m>', '', text)
    text = re.sub(r'<e_m\b[^>]*/>', '', text)
    return text

def format_date(iso_date):
    """
    Convert an ISO formatted date string into a nicer format.
    For example: '2025-03-19T21:15:44.097Z' -> 'March 19, 2025 09:15:44 PM'
    """
    try:
        if iso_date.endswith("Z"):
            iso_date = iso_date[:-1]
        dt = datetime.fromisoformat(iso_date)
        return dt.strftime("%B %d, %Y %I:%M:%S %p")
    except Exception:
        return iso_date

def filter_messages(data, username, first_word=None, conversation_id=None, search_method="first_word"):
    """
    Filter messages from the Skype conversation data and sort them by date (oldest to newest).
    """
    filtered = []
    conversations = data.get("conversations", [])
    
    for convo in conversations:
        if conversation_id and convo.get("id") != conversation_id:
            continue
        messages = convo.get("MessageList", [])
        for message in messages:
            sender = message.get("from", "")
            if ":" in sender:
                sender = sender.split(":", 1)[1]
            content = message.get("content", "")
            content = clean_content(content)
            words = content.strip().split()
            if not words:
                continue

            if sender != username:
                continue

            iso_date = message.get("originalarrivaltime", "Unknown Date")
            try:
                dt = datetime.fromisoformat(iso_date.rstrip("Z"))  # Remove 'Z' if present and parse
            except ValueError:
                dt = None  # If date is invalid, leave as None

            formatted_date = format_date(iso_date)
            
            if search_method == "first_word":
                if first_word and words[0] == first_word:
                    filtered.append({
                        "datetime": dt,
                        "date": formatted_date,
                        "content": content
                    })
            elif search_method == "review":
                if re.search(r'\b\d+(\.\d+)?/10\b', content):
                    filtered.append({
                        "datetime": dt,
                        "date": formatted_date,
                        "content": content
                    })

    # Sort messages by datetime (oldest to newest), handling None dates as well
    filtered.sort(key=lambda x: x["datetime"] if x["datetime"] else datetime.min)
    
    return filtered


def get_available_filename(filename):
    """
    Check if the file already exists. If it does, append an incrementing number to the filename.
    For example, if "user_messages.docx" exists, it will return "user_messages_1.docx", and so on.
    """
    if not os.path.exists(filename):
        return filename
    base, ext = os.path.splitext(filename)
    counter = 1
    new_filename = f"{base}_{counter}{ext}"
    while os.path.exists(new_filename):
        counter += 1
        new_filename = f"{base}_{counter}{ext}"
    return new_filename

def save_to_word(filtered_messages, output_file, username, search_method):
    """
    Save filtered messages to a Word document.
    
    The document includes a header with the username and search criteria,
    and each message is separated by a heading that includes the formatted message date.
    """
    doc = Document()
    method_text = "First word search" if search_method == "first_word" else "Video game review search"
    doc.add_heading(f"Reviews by {username}", level=1)
    
    for i, msg in enumerate(filtered_messages, start=1):
        doc.add_heading(f"Review #{i} - {msg['date']}", level=2)
        doc.add_paragraph("Message content:")
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("-" * 40)
    
    # Ensure we don't overwrite an existing file.
    final_output_file = get_available_filename(output_file)
    doc.save(final_output_file)
    print(f"Document saved as {final_output_file}")

def main():
    json_file = "skype_conversation.json"  # Update the path if necessary.
    try:
        data = load_conversation(json_file)
    except Exception as e:
        print(f"Error reading JSON file: {e}")
        return

    username = input("Enter the username to search for: ").strip()
    conversation_id = input("Enter the conversation id to filter by (or press Enter to include all): ").strip()
    if not conversation_id:
        conversation_id = None

    print("Select search method:")
    print("1 - Filter by first word in the message")
    print("2 - Filter messages that appear to be video game reviews")
    method_choice = input("Enter 1 or 2: ").strip()
    
    if method_choice == "1":
        search_method = "first_word"
        first_word = input("Enter the first word to match: ").strip()
    elif method_choice == "2":
        search_method = "review"
        first_word = None
    else:
        print("Invalid selection.")
        return

    filtered_messages = filter_messages(data, username, first_word, conversation_id, search_method)
    if not filtered_messages:
        print("No messages matched the given criteria.")
        return

    output_file = f"{username}_messages.docx"
    save_to_word(filtered_messages, output_file, username, search_method)

if __name__ == "__main__":
    main()
